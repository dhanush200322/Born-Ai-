from langchain_community.document_loaders import WebBaseLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from typing import List, Optional
import os
import pdfplumber
import docx
import logging

logger = logging.getLogger(__name__)

def parse_pdf(file_path: str) -> str:
    """Extract text from a PDF file using pdfplumber."""
    text = ""
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                content = page.extract_text()
                if content:
                    text += content + "\n"
    except Exception as e:
        logger.error(f"Error parsing PDF {file_path}: {e}")
    return text

def parse_docx(file_path: str) -> str:
    """Extract text from a DOCX file using python-docx."""
    try:
        doc = docx.Document(file_path)
        text = []
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
        return "\n".join(text)
    except Exception as e:
        logger.error(f"Error parsing DOCX {file_path}: {e}")
        return ""

def parse_txt(file_path: str) -> str:
    """Extract text from a TXT file."""
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception as e:
        logger.error(f"Error parsing TXT {file_path}: {e}")
        return ""

def process_files(file_paths: List[str]) -> List[Document]:
    """Process a list of local file paths and return langchain Documents."""
    documents = []
    for path in file_paths:
        if not os.path.exists(path):
            logger.warning(f"File not found: {path}")
            continue
            
        ext = os.path.splitext(path)[1].lower()
        text = ""
        
        if ext == ".pdf":
            text = parse_pdf(path)
        elif ext == ".docx":
            text = parse_docx(path)
        elif ext in [".txt", ".md"]:
            text = parse_txt(path)
        else:
            logger.warning(f"Unsupported file format: {ext} for path {path}")
            
        if text.strip():
            basename = os.path.basename(path)
            # Infer type from filename or extension
            source_type = "docs"
            lower_name = basename.lower()
            if "resume" in lower_name or "cv" in lower_name:
                source_type = "resume"
            
            # Remove UUID prefix if exists (uuid format: 8 chars + _)
            title = basename
            if len(basename) > 9 and basename[8] == "_":
                title = basename[9:]
                
            documents.append(
                Document(
                    page_content=text,
                    metadata={
                        "source": basename,
                        "path": path,
                        "title": title,
                        "source_type": source_type,
                        "original_filename": title
                    }
                )
            )
    return documents

def process_websites(urls: List[str]) -> List[Document]:
    """Crawl web pages and return langchain Documents."""
    documents = []
    for url in urls:
        try:
            loader = WebBaseLoader(url)
            docs = loader.load()
            for doc in docs:
                doc.metadata["source"] = url
                doc.metadata["title"] = doc.metadata.get("title", url)
                doc.metadata["source_type"] = "website"
                doc.metadata["original_filename"] = url
            documents.extend(docs)
        except Exception as e:
            logger.error(f"Error crawling {url}: {e}")
    return documents

def process_faqs(faqs: List[dict]) -> List[Document]:
    """Convert manually added FAQ Q&A pairs into langchain Documents."""
    documents = []
    for faq in faqs:
        content = f"Question: {faq['question']}\nAnswer: {faq['answer']}"
        documents.append(
            Document(
                page_content=content, 
                metadata={
                    "source": "faq",
                    "title": "FAQ",
                    "source_type": "faq",
                    "original_filename": "FAQ"
                }
            )
        )
    return documents

def chunk_documents(documents: List[Document]) -> List[Document]:
    """Split documents into smaller chunks (Size: 1000, Overlap: 200)."""
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        length_function=len
    )
    chunks = text_splitter.split_documents(documents)
    return chunks

def run_ingestion_pipeline(
    files: Optional[List[str]] = None, 
    urls: Optional[List[str]] = None, 
    faqs: Optional[List[dict]] = None
) -> List[Document]:
    """Run full ingestion, file loading, web crawling, FAQ formatting, and chunking pipeline."""
    docs = []
    
    # 1. Process local files
    if files:
        file_docs = process_files(files)
        docs.extend(file_docs)
        
    # 2. Process URLs (Website Crawling)
    if urls:
        web_docs = process_websites(urls)
        docs.extend(web_docs)
        
    # 3. Process FAQs
    if faqs:
        faq_docs = process_faqs(faqs)
        docs.extend(faq_docs)
        
    # 4. Document Chunking
    chunks = chunk_documents(docs)
    return chunks

